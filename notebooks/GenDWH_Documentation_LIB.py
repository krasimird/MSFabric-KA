# ╔══════════════════════════════════════════════════════════════════════════════╗
# ║  GenDWH — Automated Documentation System  (Library Notebook)               ║
# ║  Fabric Notebook  |  InspirIT  |  March 2026                               ║
# ║                                                                            ║
# ║  Phases: 0-Discovery → 1-Metadata → 2-AI Analysis → 3-DocGen → 4-Version  ║
# ║  Output: Word (.docx) + Excel (.xlsx) + JSONL → OneLake Documentation LH   ║
# ╚══════════════════════════════════════════════════════════════════════════════╝
#
# This notebook is a shared library loaded by each phase notebook via %run.
# It contains ALL configuration, helpers, and phase functions.
# No cell executes anything — definitions only (except the final print).

# CELL 0 ── Dependencies ───
# Install required packages (idempotent in Fabric)
# %pip install python-docx openpyxl anthropic --quiet

# CELL 1 ── Imports, CONFIG, PATHS & LIB_VERSION ───
import requests, json, time, hashlib, os, re, io, base64, shutil
from datetime import datetime, timezone
from typing import Optional
from pathlib import Path
from collections import Counter

from docx import Document as DocxDocument
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side, NamedStyle
from openpyxl.utils import get_column_letter

LIB_VERSION = "1.0.0"

CONFIG = {
    # ── Key Vault (for API keys) ──────────────────────────────────────────────
    "key_vault_name":       "kv-gendwh",
    "api_key_secret_name":  "anthropic-api-key",

    # ── Claude API ────────────────────────────────────────────────────────────
    "claude_model":         "claude-sonnet-4-20250514",
    "claude_max_tokens":    16384,

    # ── Fabric REST API ───────────────────────────────────────────────────────
    "fabric_api_base":      "https://api.fabric.microsoft.com/v1",

    # ── Documentation Lakehouse ───────────────────────────────────────────────
    "doc_lakehouse_files":  "/lakehouse/default/Files",
    "doc_lakehouse_tables": "/lakehouse/default/Tables",
    "output_dir":           "/lakehouse/default/Files/documentation",

    # ── Administration Lakehouse (metadata source) ────────────────────────────
    "admin_lakehouse_name": "GenDWH_Administration_LH",

    # ── Environment Detection ─────────────────────────────────────────────────
    "env_map": {"_WS_D": "Dev", "_WS_T": "Test", "_WS_P": "Prod"},
    "target_env":           "Dev",

    # ── Run settings ──────────────────────────────────────────────────────────
    "skip_unchanged":       True,
    "max_retries":          3,
    "retry_delay":          5,
}

PATHS = {
    "output":      CONFIG["output_dir"],
    "cache":       f"{CONFIG['doc_lakehouse_files']}/cache",
    "versions":    f"{CONFIG['doc_lakehouse_files']}/documentation/versions",
    "word_docs":   f"{CONFIG['output_dir']}/word",
    "excel_docs":  f"{CONFIG['output_dir']}/excel",
}


def _ensure_paths():
    """Create output directories (called at run-time, not at import)."""
    for p in PATHS.values():
        os.makedirs(p, exist_ok=True)


# CELL 2 ── Brand & Style Classes ───

# ── InspirIT Brand ────────────────────────────────────────────────────────────
class Brand:
    RED         = RGBColor(0xDC, 0x1F, 0x26)
    DARK        = RGBColor(0x00, 0x00, 0x00)
    GRAY_HEAD   = RGBColor(0x80, 0x80, 0x80)
    GRAY_LIGHT  = RGBColor(0xF2, 0xF3, 0xF4)
    GRAY_MID    = RGBColor(0xBD, 0xC3, 0xC7)
    GRAY_TEXT   = RGBColor(0x7F, 0x8C, 0x8D)
    WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
    TEAL        = RGBColor(0x1C, 0x8D, 0x7A)
    DARK_BLUE   = RGBColor(0x0B, 0x30, 0x52)
    GREEN       = RGBColor(0x2C, 0x6B, 0x5F)

    # Excel hex versions (no #)
    RED_HEX         = "DC1F26"
    DARK_HEX        = "000000"
    GRAY_HEAD_HEX   = "808080"
    GRAY_LIGHT_HEX  = "F2F3F4"
    GRAY_MID_HEX    = "BDC3C7"
    GRAY_TEXT_HEX   = "7F8C8D"
    WHITE_HEX       = "FFFFFF"
    TEAL_HEX        = "1C8D7A"
    DARK_BLUE_HEX   = "0B3052"
    GREEN_HEX       = "2C6B5F"
    WARM_BG_HEX     = "FFF8F0"


# ── Excel Styles ──────────────────────────────────────────────────────────────
class XS:
    """Excel style presets matching the reference lineage file."""

    THIN_BORDER = Border(
        left=Side(style='thin', color=Brand.GRAY_MID_HEX),
        right=Side(style='thin', color=Brand.GRAY_MID_HEX),
        top=Side(style='thin', color=Brand.GRAY_MID_HEX),
        bottom=Side(style='thin', color=Brand.GRAY_MID_HEX),
    )

    @staticmethod
    def title_font():
        return Font(name='Calibri Light', size=26, bold=True, color=Brand.TEAL_HEX)

    @staticmethod
    def subtitle_font():
        return Font(name='Calibri', size=10, color=Brand.GRAY_HEAD_HEX)

    @staticmethod
    def header_font():
        return Font(name='Calibri', size=10, bold=True, color=Brand.WHITE_HEX)

    @staticmethod
    def header_fill():
        return PatternFill('solid', fgColor=Brand.TEAL_HEX)

    @staticmethod
    def zone_link_font():
        return Font(name='Calibri', size=11, bold=True, color=Brand.TEAL_HEX, underline='single')

    @staticmethod
    def code_font(color=Brand.DARK_BLUE_HEX):
        return Font(name='Courier New', size=9, color=color)

    @staticmethod
    def data_font():
        return Font(name='Calibri', size=9, color=Brand.DARK_HEX)

    @staticmethod
    def meta_font():
        return Font(name='Calibri', size=9, color=Brand.GRAY_HEAD_HEX)

    @staticmethod
    def alt_fill():
        return PatternFill('solid', fgColor=Brand.GRAY_LIGHT_HEX)

    @staticmethod
    def badge_fill(transform_type):
        """Color-coded badge for transformation types."""
        mapping = {
            'direct_map':   Brand.TEAL_HEX,
            'case_when':    Brand.DARK_BLUE_HEX,
            'coalesce':     Brand.GREEN_HEX,
            'literal':      Brand.GRAY_HEAD_HEX,
            'date_format':  "5DADE2",
            'hash':         "8E44AD",
            'aggregate':    "E67E22",
            'arithmetic':   "E67E22",
            'cast':         "5DADE2",
        }
        color = mapping.get(transform_type, Brand.GRAY_HEAD_HEX)
        return PatternFill('solid', fgColor=color)

    @staticmethod
    def badge_font():
        return Font(name='Calibri', size=9, bold=True, color=Brand.WHITE_HEX)

    @staticmethod
    def nav_font():
        return Font(name='Calibri', size=9, color=Brand.TEAL_HEX, underline='single')


# ── Word (python-docx) Style Helpers — InspirIT template ─────────────────────

def docx_setup_styles(doc):
    """Configure InspirIT styles on a python-docx Document."""
    style = doc.styles
    style['Normal'].font.name = 'Calibri'
    style['Normal'].font.size = Pt(11)
    style['Normal'].font.color.rgb = Brand.DARK

    h1 = style['Heading 1']
    h1.font.name = 'Calibri Light'
    h1.font.size = Pt(28)
    h1.font.bold = True
    h1.font.color.rgb = Brand.RED
    h1.paragraph_format.space_before = Pt(16)
    h1.paragraph_format.space_after = Pt(5)

    h2 = style['Heading 2']
    h2.font.name = 'Calibri Light'
    h2.font.size = Pt(22)
    h2.font.bold = True
    h2.font.color.rgb = Brand.GRAY_HEAD
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(4)

    h3 = style['Heading 3']
    h3.font.name = 'Calibri Light'
    h3.font.size = Pt(16)
    h3.font.bold = True
    h3.font.color.rgb = Brand.GRAY_HEAD
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(3)

    return doc


def docx_add_header_footer(doc, header_text="InspirIT — Technical Documentation"):
    """Add InspirIT header/footer to all sections."""
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

        header = section.header
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = hp.add_run(header_text)
        run.font.size = Pt(9)
        run.font.color.rgb = Brand.GRAY_TEXT
        run.font.name = 'Calibri'
        pPr = hp._p.get_or_add_pPr()
        pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="4" w:space="4" w:color="DC1F26"/></w:pBdr>')
        pPr.append(pBdr)

        footer = section.footer
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run("InspirIT — The Data Platform Company")
        run.font.size = Pt(9)
        run.font.color.rgb = Brand.GRAY_TEXT
        run.font.name = 'Calibri'
        pPr = fp._p.get_or_add_pPr()
        pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:top w:val="single" w:sz="4" w:space="4" w:color="DC1F26"/></w:pBdr>')
        pPr.append(pBdr)


def docx_title_page(doc, product, title, subtitle, meta_line, date_str, meta_rows=None):
    """Add InspirIT-styled title page."""
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(product)
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.name = 'Calibri Light'
    run.font.color.rgb = Brand.RED

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.name = 'Calibri Light'
    run.font.color.rgb = Brand.DARK

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.font.size = Pt(13)
    run.font.name = 'Calibri'
    run.font.color.rgb = Brand.GRAY_TEXT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(meta_line)
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    run.font.color.rgb = Brand.GRAY_TEXT
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="6" w:space="1" w:color="DC1F26"/>'
        f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="DC1F26"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

    if meta_rows:
        doc.add_paragraph()
        for label, value in meta_rows:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"{label}:  ")
            run.font.size = Pt(10)
            run.font.color.rgb = Brand.GRAY_TEXT
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(value)
            run.font.size = Pt(10)
            run.font.name = 'Courier New'
            run.font.color.rgb = Brand.RED

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(date_str)
    run.font.size = Pt(10)
    run.font.color.rgb = Brand.GRAY_TEXT

    doc.add_page_break()


def docx_add_table(doc, headers, rows, col_widths_cm=None):
    """Add a styled table to the document."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = Brand.WHITE
        run.font.name = 'Calibri'
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="808080" w:val="clear"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val) if val else '')
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
            if r_idx % 2 == 0:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F3F4" w:val="clear"/>')
                cell._tc.get_or_add_tcPr().append(shading)

    if col_widths_cm:
        for i, w in enumerate(col_widths_cm):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table


def docx_version_table(doc, versions):
    """Add a Document History table."""
    doc.add_heading('Document History', level=2)
    headers = ['Version', 'Date', 'Author', 'Changes']
    rows = [[v['version'], v['date'], v['author'], v['description']] for v in versions]
    docx_add_table(doc, headers, rows, col_widths_cm=[2, 2.5, 3, 10])


# CELL 3 ── API Helpers ───

def get_fabric_token():
    """Get Fabric REST API bearer token via mssparkutils."""
    return mssparkutils.credentials.getToken("pbi")


def get_claude_api_key():
    """Get Anthropic API key from Key Vault."""
    return mssparkutils.credentials.getSecret(
        CONFIG["key_vault_name"],
        CONFIG["api_key_secret_name"]
    )


def fabric_api_get(endpoint, token=None):
    """GET request to Fabric REST API with auto-pagination."""
    if token is None:
        token = get_fabric_token()
    headers = {"Authorization": f"Bearer {token}"}
    url = f"{CONFIG['fabric_api_base']}{endpoint}"
    all_items = []

    while url:
        resp = requests.get(url, headers=headers)
        resp.raise_for_status()
        data = resp.json()
        all_items.extend(data.get("value", []))
        url = data.get("continuationUri") or data.get("@odata.nextLink")

    return all_items


def fabric_api_post(endpoint, body=None, token=None):
    """POST request to Fabric REST API."""
    if token is None:
        token = get_fabric_token()
    headers = {"Authorization": f"Bearer {token}", "Content-Type": "application/json"}
    url = f"{CONFIG['fabric_api_base']}{endpoint}"
    resp = requests.post(url, headers=headers, json=body or {})
    resp.raise_for_status()
    return resp.json() if resp.content else {}


def call_claude(prompt, system_prompt="", api_key=None):
    """Call Claude API with retry logic."""
    if api_key is None:
        api_key = get_claude_api_key()

    headers = {
        "x-api-key": api_key,
        "content-type": "application/json",
        "anthropic-version": "2023-06-01",
    }
    body = {
        "model": CONFIG["claude_model"],
        "max_tokens": CONFIG["claude_max_tokens"],
        "messages": [{"role": "user", "content": prompt}],
    }
    if system_prompt:
        body["system"] = system_prompt

    for attempt in range(CONFIG["max_retries"]):
        try:
            resp = requests.post(
                "https://api.anthropic.com/v1/messages",
                headers=headers, json=body, timeout=120
            )
            if resp.status_code == 429:
                wait = CONFIG["retry_delay"] * (attempt + 1)
                print(f"  ⏳ Rate limited, waiting {wait}s...")
                time.sleep(wait)
                continue
            resp.raise_for_status()
            data = resp.json()
            return data["content"][0]["text"]
        except Exception as e:
            if attempt == CONFIG["max_retries"] - 1:
                raise
            print(f"  ⚠ Attempt {attempt+1} failed: {e}, retrying...")
            time.sleep(CONFIG["retry_delay"])

    return None


def sha256(text):
    """Quick hash for change detection."""
    return hashlib.sha256(text.encode('utf-8')).hexdigest()[:16]


# CELL 4 ── Phase 0 — Discovery ───

def phase_0_discovery():
    """
    Phase 0 — Discover all Fabric workspaces and their items.
    Returns a list of dicts: [{workspace_id, workspace_name, item_id, item_type, ...}]
    """
    print("═" * 60)
    print("PHASE 0 — DISCOVERY")
    print("═" * 60)

    token = get_fabric_token()
    run_ts = datetime.now(timezone.utc).isoformat()

    # ── Step 1: List all workspaces ───────────────────────────────────────────
    print("\n📂 Listing workspaces...")
    workspaces = fabric_api_get("/workspaces", token)
    print(f"   Found {len(workspaces)} workspace(s)")

    # ── Step 2: For each workspace, list all items ────────────────────────────
    inventory = []
    for ws in workspaces:
        ws_id   = ws["id"]
        ws_name = ws.get("displayName", ws_id)
        ws_desc = ws.get("description", "")
        ws_type = ws.get("type", "")
        ws_cap  = ws.get("capacityId", "")

        print(f"\n   📁 {ws_name} ({ws_id[:8]}...)")

        try:
            items = fabric_api_get(f"/workspaces/{ws_id}/items", token)
            print(f"      {len(items)} items")
        except Exception as e:
            print(f"      ⚠ Error listing items: {e}")
            items = []

        for item in items:
            inventory.append({
                "discovery_run_ts":   run_ts,
                "workspace_id":       ws_id,
                "workspace_name":     ws_name,
                "workspace_desc":     ws_desc,
                "workspace_type":     ws_type,
                "capacity_id":        ws_cap,
                "item_id":            item.get("id", ""),
                "item_type":          item.get("type", ""),
                "item_name":          item.get("displayName", ""),
                "item_description":   item.get("description", ""),
            })

    # ── Step 3: Save to Delta table ───────────────────────────────────────────
    print(f"\n💾 Saving {len(inventory)} items to doc_workspace_inventory...")

    df = spark.createDataFrame(inventory)
    df.write.format("delta").mode("overwrite").saveAsTable("doc_workspace_inventory")

    # ── Summary ───────────────────────────────────────────────────────────────
    print("\n── Discovery Summary ──")
    ws_count = len(set(r["workspace_id"] for r in inventory))
    type_counts = {}
    for r in inventory:
        t = r["item_type"]
        type_counts[t] = type_counts.get(t, 0) + 1

    print(f"   Workspaces:  {ws_count}")
    print(f"   Total items: {len(inventory)}")
    for t, c in sorted(type_counts.items(), key=lambda x: -x[1]):
        print(f"     {t:30s} {c}")

    print("\n✓ Phase 0 complete")
    return inventory



# CELL 5 ── Phase 1 — Metadata Extraction ───

def _extract_pipeline_definition(ws_id, item_id, token):
    """Extract pipeline activities and dependencies from definition JSON."""
    try:
        result = fabric_api_post(
            f"/workspaces/{ws_id}/items/{item_id}/getDefinition", token=token
        )
        if not result:
            return None

        parts = result.get("definition", {}).get("parts", [])
        for part in parts:
            if part.get("path", "").endswith(".json"):
                payload = part.get("payload", "")
                decoded = base64.b64decode(payload).decode("utf-8")
                return json.loads(decoded)
        return None
    except Exception as e:
        print(f"      ⚠ Pipeline def error: {e}")
        return None


def _extract_pipeline_activities(pipeline_json):
    """Parse pipeline JSON into activity records."""
    activities = []
    if not pipeline_json:
        return activities

    props = pipeline_json.get("properties", {})
    for act in props.get("activities", []):
        deps = []
        for dep in act.get("dependsOn", []):
            deps.append({
                "activity": dep.get("activity", ""),
                "conditions": list(dep.get("dependencyConditions", []))
            })

        activities.append({
            "activity_name":  act.get("name", ""),
            "activity_type":  act.get("type", ""),
            "description":    act.get("description", ""),
            "state":          act.get("state", "Active"),
            "depends_on":     json.dumps(deps),
            "policy":         json.dumps(act.get("policy", {})),
            "type_props":     json.dumps(act.get("typeProperties", {}))[:4000],
        })
    return activities


def _extract_lakehouse_tables(lakehouse_name):
    """Use Spark SQL to list tables and their schemas in a lakehouse."""
    tables_info = []
    try:
        tables = spark.sql(f"SHOW TABLES IN {lakehouse_name}").collect()
        for t in tables:
            tbl_name = t["tableName"]
            try:
                cols = spark.sql(f"DESCRIBE TABLE {lakehouse_name}.{tbl_name}").collect()
                schema = [{"name": c["col_name"], "type": c["data_type"]}
                          for c in cols if not c["col_name"].startswith("#")]
                row_count = spark.sql(f"SELECT COUNT(*) as cnt FROM {lakehouse_name}.{tbl_name}").first()["cnt"]
            except:
                schema = []
                row_count = -1

            tables_info.append({
                "lakehouse":  lakehouse_name,
                "table_name": tbl_name,
                "column_count": len(schema),
                "row_count":  row_count,
                "schema_json": json.dumps(schema),
            })
    except Exception as e:
        print(f"      ⚠ LH tables error ({lakehouse_name}): {e}")
    return tables_info


def _extract_metadata_queries(admin_lh_name):
    """
    Read SQL queries from gen_adm_* metadata tables.
    Returns list of {meta_table, target_table, layer, mode, level, source_query, merge_key, ...}
    """
    queries = []
    meta_tables = [
        # Silver
        ("gen_adm_scd1_merge_silver_level_1_meta",    "Silver", "merge",     "L1"),
        ("gen_adm_scd1_merge_silver_level_2_meta",    "Silver", "merge",     "L2"),
        ("gen_adm_overwrite_silver_level_1_meta",     "Silver", "overwrite", "L1"),
        ("gen_adm_overwrite_silver_level_2_meta",     "Silver", "overwrite", "L2"),
        # Gold
        ("gen_adm_scd2_merge_gold_level_1_meta",      "Gold",   "merge",     "L1"),
        ("gen_adm_scd2_merge_gold_level_2_meta",      "Gold",   "merge",     "L2"),
        ("gen_adm_scd2_merge_gold_level_3_meta",      "Gold",   "merge",     "L3"),
        ("gen_adm_scd2_merge_gold_level_4_meta",      "Gold",   "merge",     "L4"),
        ("gen_adm_scd2_merge_gold_level_5_meta",      "Gold",   "merge",     "L5"),
        ("gen_adm_overwrite_gold_level_1_meta",       "Gold",   "overwrite", "L1"),
        ("gen_adm_overwrite_gold_level_2_meta",       "Gold",   "overwrite", "L2"),
    ]

    for meta_tbl, layer, mode, level in meta_tables:
        try:
            rows = spark.sql(f"SELECT * FROM {admin_lh_name}.{meta_tbl}").collect()
            for row in rows:
                rd = row.asDict()
                source_query = rd.get("source_query", "") or ""
                target_table = rd.get("target_table_name", "") or rd.get("target_table", "") or ""
                merge_key    = rd.get("merge_key", "") or rd.get("business_key", "") or ""

                queries.append({
                    "meta_table":    meta_tbl,
                    "target_table":  target_table,
                    "layer":         layer,
                    "mode":          mode,
                    "level":         level,
                    "source_query":  source_query,
                    "merge_key":     merge_key,
                    "query_hash":    sha256(source_query) if source_query else "",
                    "has_current":   str(rd.get("has_current", "true")),
                    "is_active":     str(rd.get("is_active", "true")),
                })
        except Exception as e:
            print(f"   ⚠ Cannot read {meta_tbl}: {e}")

    return queries


def phase_1_metadata_extraction(inventory=None):
    """
    Phase 1 — Extract detailed metadata from all discovered items.
    """
    print("═" * 60)
    print("PHASE 1 — METADATA EXTRACTION")
    print("═" * 60)

    token = get_fabric_token()

    if inventory is None:
        inventory = spark.sql("SELECT * FROM doc_workspace_inventory").collect()
        inventory = [r.asDict() for r in inventory]

    # ── 1a: Pipeline definitions ──────────────────────────────────────────────
    print("\n🔧 Extracting pipeline definitions...")
    pipelines = [i for i in inventory if i["item_type"] == "DataPipeline"]
    all_activities = []

    for p in pipelines:
        print(f"   📋 {p['item_name']}")
        pdef = _extract_pipeline_definition(p["workspace_id"], p["item_id"], token)
        acts = _extract_pipeline_activities(pdef)
        for a in acts:
            a["workspace_id"]   = p["workspace_id"]
            a["workspace_name"] = p["workspace_name"]
            a["pipeline_id"]    = p["item_id"]
            a["pipeline_name"]  = p["item_name"]
        all_activities.extend(acts)
        print(f"      {len(acts)} activities")

    if all_activities:
        df = spark.createDataFrame(all_activities)
        df.write.format("delta").mode("overwrite").saveAsTable("doc_pipeline_activities")
        print(f"   💾 Saved {len(all_activities)} activities to doc_pipeline_activities")

    # ── 1b: Lakehouse table schemas ───────────────────────────────────────────
    print("\n🗄️  Extracting lakehouse schemas...")
    lakehouses = [i for i in inventory if i["item_type"] == "Lakehouse"]
    all_tables = []

    for lh in lakehouses:
        lh_name = lh["item_name"]
        print(f"   📦 {lh_name}")
        tables = _extract_lakehouse_tables(lh_name)
        for t in tables:
            t["workspace_id"]   = lh["workspace_id"]
            t["workspace_name"] = lh["workspace_name"]
            t["item_id"]        = lh["item_id"]
        all_tables.extend(tables)
        print(f"      {len(tables)} tables")

    if all_tables:
        df = spark.createDataFrame(all_tables)
        df.write.format("delta").mode("overwrite").saveAsTable("doc_lakehouse_tables")
        print(f"   💾 Saved {len(all_tables)} table schemas to doc_lakehouse_tables")

    # ── 1c: Metadata SQL queries (for lineage) ───────────────────────────────
    print(f"\n📜 Extracting SQL queries from {CONFIG['admin_lakehouse_name']}...")
    queries = _extract_metadata_queries(CONFIG["admin_lakehouse_name"])
    print(f"   Found {len(queries)} SQL queries")

    if queries:
        df = spark.createDataFrame(queries)
        df.write.format("delta").mode("overwrite").saveAsTable("doc_metadata_queries")
        print(f"   💾 Saved to doc_metadata_queries")

    # ── 1d: Bronze metadata ───────────────────────────────────────────────────
    print("\n🥉 Extracting Bronze metadata...")
    bronze_meta = []
    try:
        rows = spark.sql(
            f"SELECT * FROM {CONFIG['admin_lakehouse_name']}.gen_adm_meta_bronze"
        ).collect()
        for row in rows:
            rd = row.asDict()
            bronze_meta.append({
                "source_schema":  rd.get("source_schema", ""),
                "source_table":   rd.get("source_table", ""),
                "source_columns": rd.get("source_columns", ""),
                "target_table":   rd.get("target_table", ""),
                "is_active":      str(rd.get("is_active", True)),
            })
        if bronze_meta:
            df = spark.createDataFrame(bronze_meta)
            df.write.format("delta").mode("overwrite").saveAsTable("doc_bronze_metadata")
            print(f"   💾 {len(bronze_meta)} Bronze tables saved to doc_bronze_metadata")
    except Exception as e:
        print(f"   ⚠ Bronze metadata error: {e}")

    # ── 1e: Warehouse objects ─────────────────────────────────────────────────
    warehouses = [i for i in inventory if i["item_type"] == "Warehouse"]
    all_wh_objects = []
    if warehouses:
        print(f"\n🏭 Extracting warehouse objects ({len(warehouses)} warehouses)...")
        for wh in warehouses:
            print(f"   📦 {wh['item_name']}")
            objs = extract_warehouse(wh["workspace_id"], wh["item_id"], wh["item_name"], token)
            all_wh_objects.extend(objs)
            tbl_c = sum(1 for o in objs if o["object_type"] == "table")
            view_c = sum(1 for o in objs if o["object_type"] == "view")
            sproc_c = sum(1 for o in objs if o["object_type"] == "sproc")
            print(f"      {tbl_c} tables, {view_c} views, {sproc_c} sprocs")
        if all_wh_objects:
            df = spark.createDataFrame(all_wh_objects)
            df.write.format("delta").mode("overwrite").saveAsTable("doc_warehouse_objects")
            print(f"   💾 Saved {len(all_wh_objects)} objects to doc_warehouse_objects")

    # ── 1f: Notebook definitions ──────────────────────────────────────────────
    notebooks = [i for i in inventory if i["item_type"] == "Notebook"]
    all_nb_records = []
    if notebooks:
        print(f"\n📓 Extracting notebook definitions ({len(notebooks)} notebooks)...")
        for nb in notebooks:
            print(f"   📝 {nb['item_name']}")
            recs = extract_notebooks(nb["workspace_id"], nb["item_id"], nb["item_name"], token)
            for r in recs:
                r["workspace_name"] = nb["workspace_name"]
            all_nb_records.extend(recs)
        if all_nb_records:
            df = spark.createDataFrame(all_nb_records)
            df.write.format("delta").mode("overwrite").saveAsTable("doc_notebook_definitions")
            print(f"   💾 Saved {len(all_nb_records)} notebooks to doc_notebook_definitions")

    # ── 1g: Semantic models ───────────────────────────────────────────────────
    sem_models = [i for i in inventory if i["item_type"] == "SemanticModel"]
    all_sm_records = []
    if sem_models:
        print(f"\n📊 Extracting semantic models ({len(sem_models)} models)...")
        for sm in sem_models:
            print(f"   📈 {sm['item_name']}")
            recs = extract_semantic_models(sm["workspace_id"], sm["item_id"], sm["item_name"], token)
            all_sm_records.extend(recs)
            m_c = sum(1 for r in recs if r["object_type"] == "measure")
            r_c = sum(1 for r in recs if r["object_type"] == "relationship")
            print(f"      {m_c} measures, {r_c} relationships")
        if all_sm_records:
            df = spark.createDataFrame(all_sm_records)
            df.write.format("delta").mode("overwrite").saveAsTable("doc_semantic_models")
            print(f"   💾 Saved {len(all_sm_records)} objects to doc_semantic_models")

    # ── 1h: Report definitions ────────────────────────────────────────────────
    reports = [i for i in inventory
               if i["item_type"] in ("Report", "PaginatedReport")]
    all_rpt_records = []
    if reports:
        print(f"\n📄 Extracting report definitions ({len(reports)} reports)...")
        for rpt in reports:
            print(f"   📋 {rpt['item_name']}")
            recs = extract_reports(rpt["workspace_id"], rpt["item_id"], rpt["item_name"], token)
            all_rpt_records.extend(recs)
        if all_rpt_records:
            df = spark.createDataFrame(all_rpt_records)
            df.write.format("delta").mode("overwrite").saveAsTable("doc_report_definitions")
            print(f"   💾 Saved {len(all_rpt_records)} reports to doc_report_definitions")

    # ── 1i: Variable libraries ────────────────────────────────────────────────
    var_libs = [i for i in inventory if i["item_type"] == "VariableLibrary"]
    all_vl_records = []
    if var_libs:
        print(f"\n📚 Extracting variable libraries ({len(var_libs)} libraries)...")
        for vl in var_libs:
            print(f"   📖 {vl['item_name']}")
            recs = extract_variable_library(vl["workspace_id"], vl["item_id"], vl["item_name"], token)
            all_vl_records.extend(recs)
        if all_vl_records:
            df = spark.createDataFrame(all_vl_records)
            df.write.format("delta").mode("overwrite").saveAsTable("doc_variable_library")
            print(f"   💾 Saved {len(all_vl_records)} variables to doc_variable_library")

    print("\n✓ Phase 1 complete")
    return {
        "activities":      all_activities,
        "tables":          all_tables,
        "queries":         queries,
        "bronze_meta":     bronze_meta,
        "warehouse_objects": all_wh_objects,
        "notebooks":       all_nb_records,
        "semantic_models": all_sm_records,
        "reports":         all_rpt_records,
        "variable_library": all_vl_records,
    }


# ── Item-type extractors ──────────────────────────────────────────────────────

def extract_warehouse(ws_id, item_id, item_name, token):
    """
    Extract Warehouse metadata: tables, views (with SQL definitions), stored procedures.
    Uses Fabric SQL Analytics endpoint via Spark SQL to query INFORMATION_SCHEMA.
    Returns list of dicts with schema matching doc_warehouse_objects.
    """
    objects = []
    wh_name = item_name

    # Tables
    try:
        rows = spark.sql(
            f"SELECT TABLE_SCHEMA, TABLE_NAME FROM {wh_name}.INFORMATION_SCHEMA.TABLES "
            f"WHERE TABLE_TYPE = 'BASE TABLE'"
        ).collect()
        for r in rows:
            schema_json = []
            try:
                cols = spark.sql(
                    f"SELECT COLUMN_NAME, DATA_TYPE, IS_NULLABLE, ORDINAL_POSITION "
                    f"FROM {wh_name}.INFORMATION_SCHEMA.COLUMNS "
                    f"WHERE TABLE_SCHEMA = '{r['TABLE_SCHEMA']}' AND TABLE_NAME = '{r['TABLE_NAME']}' "
                    f"ORDER BY ORDINAL_POSITION"
                ).collect()
                schema_json = [
                    {"name": c["COLUMN_NAME"], "type": c["DATA_TYPE"],
                     "nullable": c["IS_NULLABLE"], "ordinal": c["ORDINAL_POSITION"]}
                    for c in cols
                ]
            except Exception:
                pass
            objects.append({
                "workspace_id": ws_id, "warehouse": wh_name,
                "object_name": f"{r['TABLE_SCHEMA']}.{r['TABLE_NAME']}",
                "object_type": "table", "definition_sql": "",
                "schema_json": json.dumps(schema_json),
            })
    except Exception as e:
        print(f"      ⚠ Warehouse tables error: {e}")

    # Views (with definition SQL)
    try:
        rows = spark.sql(
            f"SELECT TABLE_SCHEMA, TABLE_NAME, VIEW_DEFINITION "
            f"FROM {wh_name}.INFORMATION_SCHEMA.VIEWS"
        ).collect()
        for r in rows:
            objects.append({
                "workspace_id": ws_id, "warehouse": wh_name,
                "object_name": f"{r['TABLE_SCHEMA']}.{r['TABLE_NAME']}",
                "object_type": "view",
                "definition_sql": (r["VIEW_DEFINITION"] or "")[:8000],
                "schema_json": "",
            })
    except Exception as e:
        print(f"      ⚠ Warehouse views error: {e}")

    # Stored procedures
    try:
        rows = spark.sql(
            f"SELECT ROUTINE_SCHEMA, ROUTINE_NAME, ROUTINE_DEFINITION "
            f"FROM {wh_name}.INFORMATION_SCHEMA.ROUTINES "
            f"WHERE ROUTINE_TYPE = 'PROCEDURE'"
        ).collect()
        for r in rows:
            objects.append({
                "workspace_id": ws_id, "warehouse": wh_name,
                "object_name": f"{r['ROUTINE_SCHEMA']}.{r['ROUTINE_NAME']}",
                "object_type": "sproc",
                "definition_sql": (r["ROUTINE_DEFINITION"] or "")[:8000],
                "schema_json": "",
            })
    except Exception as e:
        print(f"      ⚠ Warehouse sprocs error: {e}")

    return objects


def extract_notebooks(ws_id, item_id, item_name, token):
    """
    Extract Notebook content via Fabric REST API getDefinition.
    Decodes Base64 payload → parses .ipynb JSON → extracts cells.
    Returns list of dicts with schema matching doc_notebook_definitions.
    """
    records = []
    try:
        result = fabric_api_post(
            f"/workspaces/{ws_id}/items/{item_id}/getDefinition", token=token
        )
        if not result:
            return records

        parts = result.get("definition", {}).get("parts", [])
        nb_json = None
        for part in parts:
            if part.get("path", "").endswith(".ipynb"):
                payload = part.get("payload", "")
                decoded = base64.b64decode(payload).decode("utf-8")
                nb_json = json.loads(decoded)
                break

        if not nb_json:
            return records

        cells = nb_json.get("cells", [])
        code_cells = [c for c in cells if c.get("cell_type") == "code"]
        md_cells = [c for c in cells if c.get("cell_type") == "markdown"]

        # Extract notebook-level parameters from metadata
        nb_meta = nb_json.get("metadata", {})
        params = nb_meta.get("widgets", {}).get("application/vnd.databricks.v1+cell", {})
        if not params:
            params = nb_meta.get("parameters", {})

        source_code = "\n".join(
            "".join(c.get("source", [])) for c in code_cells
        )
        markdown_content = "\n".join(
            "".join(c.get("source", [])) for c in md_cells
        )

        records.append({
            "workspace_id":   ws_id,
            "workspace_name": "",   # filled by caller
            "notebook_name":  item_name,
            "parameters":     json.dumps(params)[:4000],
            "cell_count":     len(cells),
            "source_code":    source_code[:32000],
            "markdown_cells": markdown_content[:16000],
        })
    except Exception as e:
        print(f"      ⚠ Notebook extraction error: {e}")

    return records


def extract_semantic_models(ws_id, item_id, item_name, token):
    """
    Extract Semantic Model definition via Fabric REST API getDefinition.
    Parses TMDL/TMSL for measures (DAX), relationships, and calculated columns.
    Returns list of dicts with schema matching doc_semantic_models.
    """
    records = []
    try:
        result = fabric_api_post(
            f"/workspaces/{ws_id}/items/{item_id}/getDefinition",
            body={"format": "TMDL"},
            token=token,
        )
        if not result:
            return records

        parts = result.get("definition", {}).get("parts", [])

        for part in parts:
            path = part.get("path", "")
            payload = part.get("payload", "")
            try:
                content = base64.b64decode(payload).decode("utf-8")
            except Exception:
                continue

            # Parse .tmdl table files for measures and calculated columns
            if "/tables/" in path.lower() and path.endswith(".tmdl"):
                table_name = path.split("/")[-1].replace(".tmdl", "")
                _parse_tmdl_table(records, ws_id, item_name, table_name, content)

            # Parse relationships file
            if path.lower().endswith("relationships.tmdl") or "/relationships/" in path.lower():
                _parse_tmdl_relationships(records, ws_id, item_name, content)

    except Exception as e:
        print(f"      ⚠ Semantic model extraction error: {e}")

    return records


def _parse_tmdl_table(records, ws_id, model_name, table_name, content):
    """Parse a TMDL table definition for measures and calculated columns."""
    lines = content.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].strip()

        # Detect measure blocks
        if line.startswith("measure "):
            measure_name = line.replace("measure ", "").strip().rstrip("=").strip()
            expression_lines = []
            format_string = ""
            i += 1
            while i < len(lines) and not (lines[i].strip().startswith("measure ") and "=" in lines[i]):
                l = lines[i].strip()
                if l.startswith("formatString:"):
                    format_string = l.split(":", 1)[1].strip().strip('"').strip("'")
                elif not l.startswith("annotation ") and not l.startswith("displayFolder"):
                    if l and not l.startswith("lineageTag:"):
                        expression_lines.append(l)
                # Break if we hit next top-level block
                if lines[i] and not lines[i][0].isspace() and i > 0:
                    break
                i += 1
            records.append({
                "workspace_id": ws_id, "model_name": model_name,
                "object_type": "measure", "object_name": measure_name,
                "expression": "\n".join(expression_lines)[:4000],
                "format_string": format_string,
                "source_table": table_name, "target_table": "",
            })
            continue

        # Detect calculated columns
        if line.startswith("column ") and "=" in line:
            col_parts = line.replace("column ", "").split("=", 1)
            col_name = col_parts[0].strip()
            expression = col_parts[1].strip() if len(col_parts) > 1 else ""
            records.append({
                "workspace_id": ws_id, "model_name": model_name,
                "object_type": "calculated_column", "object_name": col_name,
                "expression": expression[:4000],
                "format_string": "", "source_table": table_name, "target_table": "",
            })

        i += 1


def _parse_tmdl_relationships(records, ws_id, model_name, content):
    """Parse TMDL relationship definitions."""
    lines = content.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if line.startswith("relationship "):
            rel_name = line.replace("relationship ", "").strip()
            from_table = from_col = to_table = to_col = ""
            i += 1
            while i < len(lines) and (lines[i].startswith(" ") or lines[i].startswith("\t")):
                l = lines[i].strip()
                if l.startswith("fromColumn:"):
                    from_col = l.split(":", 1)[1].strip()
                elif l.startswith("fromTable:"):
                    from_table = l.split(":", 1)[1].strip()
                elif l.startswith("toColumn:"):
                    to_col = l.split(":", 1)[1].strip()
                elif l.startswith("toTable:"):
                    to_table = l.split(":", 1)[1].strip()
                i += 1
            records.append({
                "workspace_id": ws_id, "model_name": model_name,
                "object_type": "relationship", "object_name": rel_name,
                "expression": f"{from_table}.{from_col} -> {to_table}.{to_col}",
                "format_string": "", "source_table": from_table, "target_table": to_table,
            })
            continue
        i += 1


def extract_reports(ws_id, item_id, item_name, token):
    """
    Extract Report definition via Fabric REST API getDefinition.
    Parses report.json for pages, semantic model binding, filters, bookmarks.
    Handles both Report (.pbir) and PaginatedReport (.rdl) types.
    Returns list of dicts with schema matching doc_report_definitions.
    """
    records = []
    try:
        result = fabric_api_post(
            f"/workspaces/{ws_id}/items/{item_id}/getDefinition", token=token
        )
        if not result:
            return records

        parts = result.get("definition", {}).get("parts", [])
        report_type = "Report"
        semantic_model = ""
        pages = []
        filters = []
        bookmarks = []

        for part in parts:
            path = part.get("path", "")
            payload = part.get("payload", "")
            try:
                content = base64.b64decode(payload).decode("utf-8")
            except Exception:
                continue

            # Detect PaginatedReport via RDL XML
            if path.lower().endswith(".rdl"):
                report_type = "PaginatedReport"
                # Extract data source from RDL XML (basic parsing)
                ds_match = re.findall(r"<DataSourceName>(.*?)</DataSourceName>", content)
                if ds_match:
                    semantic_model = ds_match[0]
                # Extract report sections/pages from RDL
                page_matches = re.findall(r"<ReportSection[^>]*>(.*?)</ReportSection>", content, re.DOTALL)
                for idx, _ in enumerate(page_matches):
                    pages.append({"ordinal": idx, "name": f"Page {idx + 1}"})
                records.append({
                    "workspace_id": ws_id, "report_name": item_name,
                    "report_type": report_type, "semantic_model": semantic_model,
                    "pages": json.dumps(pages), "filters": json.dumps(filters),
                    "bookmarks": json.dumps(bookmarks),
                })
                return records

            # Standard report definition.pbir or report.json
            if path.lower().endswith((".pbir", "report.json")):
                try:
                    rjson = json.loads(content)
                    semantic_model = (
                        rjson.get("datasetReference", {}).get("byPath", {}).get("path", "")
                        or rjson.get("datasetReference", {}).get("byConnection", {}).get("connectionString", "")
                    )
                except Exception:
                    pass

            # Parse pages
            if "/pages/" in path.lower() and path.lower().endswith(".json"):
                try:
                    pjson = json.loads(content)
                    pages.append({
                        "name": pjson.get("displayName", path.split("/")[-1]),
                        "ordinal": pjson.get("ordinal", len(pages)),
                        "visual_count": len(pjson.get("visualContainers", [])),
                    })
                except Exception:
                    pass

            # Parse filters
            if "filters" in path.lower() and path.lower().endswith(".json"):
                try:
                    fjson = json.loads(content)
                    if isinstance(fjson, list):
                        filters.extend(fjson)
                    else:
                        filters.append(fjson)
                except Exception:
                    pass

            # Parse bookmarks
            if "bookmarks" in path.lower() and path.lower().endswith(".json"):
                try:
                    bjson = json.loads(content)
                    if isinstance(bjson, list):
                        bookmarks.extend(bjson)
                    else:
                        bookmarks.append(bjson)
                except Exception:
                    pass

        records.append({
            "workspace_id": ws_id, "report_name": item_name,
            "report_type": report_type, "semantic_model": semantic_model,
            "pages": json.dumps(pages), "filters": json.dumps(filters),
            "bookmarks": json.dumps(bookmarks),
        })

    except Exception as e:
        print(f"      ⚠ Report extraction error: {e}")

    return records


def extract_variable_library(ws_id, item_id, item_name, token):
    """
    Extract Variable Library entries via Fabric REST API.
    Returns list of dicts with schema matching doc_variable_library.
    """
    records = []
    try:
        result = fabric_api_post(
            f"/workspaces/{ws_id}/items/{item_id}/getDefinition", token=token
        )
        if not result:
            return records

        parts = result.get("definition", {}).get("parts", [])
        for part in parts:
            payload = part.get("payload", "")
            try:
                content = base64.b64decode(payload).decode("utf-8")
                vlib = json.loads(content)
            except Exception:
                continue

            # Variable library JSON can be a dict with variable entries
            variables = vlib if isinstance(vlib, dict) else {}
            if isinstance(vlib, dict) and "variables" in vlib:
                variables = vlib["variables"]

            if isinstance(variables, dict):
                for var_name, var_val in variables.items():
                    v_type = type(var_val).__name__
                    if isinstance(var_val, dict):
                        v_type = var_val.get("type", "object")
                        v_value = json.dumps(var_val.get("value", var_val))
                    else:
                        v_value = str(var_val)
                    records.append({
                        "workspace_id": ws_id, "library_name": item_name,
                        "variable_name": var_name, "variable_value": v_value[:4000],
                        "variable_type": v_type,
                    })
            elif isinstance(variables, list):
                for v in variables:
                    if isinstance(v, dict):
                        records.append({
                            "workspace_id": ws_id, "library_name": item_name,
                            "variable_name": v.get("name", ""),
                            "variable_value": str(v.get("value", ""))[:4000],
                            "variable_type": v.get("type", "string"),
                        })

    except Exception as e:
        print(f"      ⚠ Variable library extraction error: {e}")

    return records


# CELL 6 ── Phase 2 — AI Analysis ───

LINEAGE_SYSTEM_PROMPT = """You are a data lineage analyst for the GenDWH platform (Generali DWH on Microsoft Fabric).
Your task: analyze a SQL query and produce a precise JSON field-level lineage.

RULES:
1. Resolve ALL CTE alias chains to the original Bronze/Silver/Gold source tables
2. For each target field, identify: source table, source column, transformation type, SQL expression, business logic
3. Transformation types: direct_map, case_when, coalesce, cast, arithmetic, hash, date_format, concat, aggregate, literal, subquery, expression
4. For literals/constants (hardcoded values like -1, '-', 'ClEx'), set source_table="—" and source_column="—"
5. Include join keys used to connect source tables
6. Mark SCD2 tracked fields if has_current is mentioned

RESPOND WITH ONLY VALID JSON — no markdown, no preamble. Schema:
{
  "target_table": "table_name",
  "source_tables": ["full.table.name", ...],
  "cte_chain": ["cte_name → source_description", ...],
  "join_keys": ["key1 = key2", ...],
  "fields": [
    {
      "target_field": "field_name",
      "data_type": "string|int|bigint|decimal|date|...",
      "source_table": "full.source.table or — for literals",
      "source_column": "column_name or — for literals",
      "transformation_type": "direct_map|case_when|coalesce|...",
      "expression": "SQL expression or empty",
      "business_logic": "Brief human-readable explanation in English",
      "join_key": "join condition if relevant, else empty"
    }
  ]
}"""


def _build_lineage_prompt(query_rec):
    """Build the user prompt for a single SQL query analysis."""
    return f"""Analyze this SQL query for the GenDWH platform.

Target table: {query_rec['target_table']}
Layer: {query_rec['layer']}
Mode: {query_rec['mode']} (Level: {query_rec['level']})
Merge key: {query_rec.get('merge_key', '')}
has_current: {query_rec.get('has_current', 'true')}

SQL Query:
```sql
{query_rec['source_query']}
```

Produce the field-level lineage JSON. Resolve all CTE chains to original source tables."""


def _load_cache():
    """Load AI analysis cache from JSON file."""
    cache_path = f"{PATHS['cache']}/ai_lineage_cache.json"
    if os.path.exists(cache_path):
        with open(cache_path, 'r') as f:
            return json.load(f)
    return {}


def _save_cache(cache):
    """Save AI analysis cache to JSON file."""
    cache_path = f"{PATHS['cache']}/ai_lineage_cache.json"
    with open(cache_path, 'w') as f:
        json.dump(cache, f, indent=2, ensure_ascii=False)


def phase_2_ai_analysis(queries=None, force_rerun=False):
    """
    Phase 2 — Send SQL queries to Claude for field-level lineage analysis.
    Caches results keyed by query_hash.
    """
    print("═" * 60)
    print("PHASE 2 — AI ANALYSIS")
    print("═" * 60)

    if queries is None:
        rows = spark.sql("SELECT * FROM doc_metadata_queries").collect()
        queries = [r.asDict() for r in rows]

    # Filter to queries with actual SQL
    actionable = [q for q in queries if q.get("source_query", "").strip()]
    print(f"\n📊 {len(actionable)} queries with SQL (out of {len(queries)} total)")

    api_key = get_claude_api_key()
    cache = _load_cache()
    results = {}
    skipped = 0
    analyzed = 0
    failed = 0

    for i, q in enumerate(actionable):
        qhash = q["query_hash"]
        tbl = q["target_table"]

        # Check cache
        if not force_rerun and CONFIG["skip_unchanged"] and qhash in cache:
            results[tbl] = cache[qhash]
            skipped += 1
            continue

        print(f"\n   [{i+1}/{len(actionable)}] 🔍 {tbl} ({q['layer']} {q['level']} {q['mode']})")

        try:
            prompt = _build_lineage_prompt(q)
            response = call_claude(prompt, LINEAGE_SYSTEM_PROMPT, api_key)

            # Parse JSON — strip markdown fences if present
            clean = response.strip()
            if clean.startswith("```"):
                clean = re.sub(r'^```\w*\n?', '', clean)
                clean = re.sub(r'\n?```$', '', clean)

            lineage = json.loads(clean)
            lineage["_meta"] = {
                "query_hash": qhash,
                "analyzed_at": datetime.now(timezone.utc).isoformat(),
                "layer": q["layer"],
                "mode": q["mode"],
                "level": q["level"],
                "merge_key": q.get("merge_key", ""),
                "has_current": q.get("has_current", "true"),
            }

            results[tbl] = lineage
            cache[qhash] = lineage
            analyzed += 1
            print(f"      ✓ {len(lineage.get('fields', []))} fields extracted")

            # Save cache periodically (every 5 queries)
            if analyzed % 5 == 0:
                _save_cache(cache)

        except json.JSONDecodeError as e:
            print(f"      ⚠ JSON parse error: {e}")
            failed += 1
        except Exception as e:
            print(f"      ⚠ Error: {e}")
            failed += 1

    # Final cache save
    _save_cache(cache)

    # Save results as Delta table
    flat_records = []
    for tbl, lineage in results.items():
        meta = lineage.get("_meta", {})
        for field in lineage.get("fields", []):
            flat_records.append({
                "target_table":       tbl,
                "layer":              meta.get("layer", ""),
                "mode":               meta.get("mode", ""),
                "level":              meta.get("level", ""),
                "target_field":       field.get("target_field", ""),
                "data_type":          field.get("data_type", ""),
                "source_table":       field.get("source_table", ""),
                "source_column":      field.get("source_column", ""),
                "transformation_type":field.get("transformation_type", ""),
                "expression":         field.get("expression", ""),
                "business_logic":     field.get("business_logic", ""),
                "join_key":           field.get("join_key", ""),
                "query_hash":         meta.get("query_hash", ""),
                "analyzed_at":        meta.get("analyzed_at", ""),
            })

    if flat_records:
        df = spark.createDataFrame(flat_records)
        df.write.format("delta").mode("overwrite").saveAsTable("doc_ai_lineage")
        print(f"\n💾 Saved {len(flat_records)} field records to doc_ai_lineage")

    print(f"\n── AI Analysis Summary ──")
    print(f"   Analyzed: {analyzed}")
    print(f"   Cached (skipped): {skipped}")
    print(f"   Failed: {failed}")
    print(f"   Total fields: {len(flat_records)}")
    print("\n✓ Phase 2 complete")

    return results



# CELL 7 ── Phase 3 — Knowledge Build (JSONL) ───

CHUNK_TYPES = [
    "table_lineage",
    "field_detail",
    "pipeline_overview",
    "pipeline_activity",
    "table_schema",
    "warehouse_view",
    "warehouse_sproc",
    "notebook_definition",
    "semantic_measure",
    "semantic_relationship",
    "report_definition",
    "paginated_report",
    "variable",
    "glossary_term",
    "architecture",
]


def phase_3_knowledge_build(ai_results=None, inventory=None, activities=None,
                            queries=None, lh_tables=None, bronze_meta=None):
    """Phase 3 — Build the unified JSONL knowledge base for RAG.

    Reads all doc_* Delta tables produced by Phases 0-2 and writes
    gendwh_knowledge.jsonl — one JSON object per line, each tagged with
    a chunk_type from CHUNK_TYPES.

    The JSONL file is consumed by the webapp's retrieval layer in Sprint 3.

    Implementation deferred to Sprint 2 (S2.1–S2.5).
    """
    pass



# CELL 8 ── Phase 4 — Doc Generation ───

def _xl_apply_header_row(ws, row_num, num_cols):
    """Apply teal header styling to a row."""
    for col in range(1, num_cols + 1):
        cell = ws.cell(row=row_num, column=col)
        cell.font = XS.header_font()
        cell.fill = XS.header_fill()
        cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        cell.border = XS.THIN_BORDER


def _xl_apply_data_row(ws, row_num, num_cols, is_alt=False):
    """Apply data row styling with alternate shading."""
    for col in range(1, num_cols + 1):
        cell = ws.cell(row=row_num, column=col)
        cell.font = XS.data_font()
        cell.alignment = Alignment(vertical='center', wrap_text=True)
        cell.border = XS.THIN_BORDER
        if is_alt:
            cell.fill = XS.alt_fill()


def _xl_safe_sheet_name(name, prefix="", max_len=31):
    """Create a safe sheet name within Excel's 31-char limit."""
    full = f"{prefix}{name}" if prefix else name
    for ch in ['\\', '/', '*', '?', ':', '[', ']']:
        full = full.replace(ch, '_')
    return full[:max_len]


def _xl_add_nav_row(ws, row, links):
    """Add navigation links in row 1: ← Zone Index | ↖ INDEX"""
    col = 1
    for label, sheet_name in links:
        cell = ws.cell(row=row, column=col)
        cell.value = label
        cell.font = XS.nav_font()
        cell.hyperlink = f"#'{sheet_name}'!A1"
        col += 1


# ── Data Lineage Excel ────────────────────────────────────────────────────────

def generate_lineage_excel(ai_results, bronze_meta=None, queries=None):
    """
    Generate Data Lineage Excel with navigation: INDEX → Zone → Detail.
    Matches the reference file structure.
    """
    print("\n📗 Generating Data Lineage Excel...")
    wb = openpyxl.Workbook()
    ws_index = wb.active
    ws_index.title = "INDEX"
    now_str = datetime.now().strftime("%B %Y")

    # ── Classify tables by zone ───────────────────────────────────────────────
    zones = {"Bronze": [], "Silver": [], "Gold": []}

    if bronze_meta:
        for bm in bronze_meta:
            zones["Bronze"].append({
                "target_table": bm.get("target_table", ""),
                "source_schema": bm.get("source_schema", ""),
                "source_table": bm.get("source_table", ""),
                "fields_count": bm.get("source_columns", "*"),
                "is_active": bm.get("is_active", "true"),
            })

    for tbl, lineage in (ai_results or {}).items():
        meta = lineage.get("_meta", {})
        layer = meta.get("layer", "")
        zone = "Silver" if layer == "Silver" else "Gold" if layer == "Gold" else None
        if zone:
            zones[zone].append({
                "target_table": tbl,
                "lineage": lineage,
                "mode": meta.get("mode", ""),
                "level": meta.get("level", ""),
                "merge_key": meta.get("merge_key", ""),
                "has_current": meta.get("has_current", "true"),
                "fields_count": len(lineage.get("fields", [])),
                "source_tables": lineage.get("source_tables", []),
            })

    zone_sheet_names = {
        "Bronze": "Bronze",
        "Silver": "SilverRaw_SilverStg",
        "Gold":   "GoldDWH",
    }

    # ── INDEX sheet ───────────────────────────────────────────────────────────
    ws_index.merge_cells('A1:C1')
    ws_index['A1'] = "GenDWH — Data Lineage"
    ws_index['A1'].font = XS.title_font()

    ws_index.merge_cells('A2:C2')
    ws_index['A2'] = f"Bronze → Silver → Gold  ·  GenDWH Platform  ·  Microsoft Fabric  ·  InspirIT  ·  {now_str}"
    ws_index['A2'].font = XS.subtitle_font()

    for col, hdr in enumerate(['Zone', 'Description', 'Tables'], 1):
        cell = ws_index.cell(row=4, column=col, value=hdr)
    _xl_apply_header_row(ws_index, 4, 3)

    zone_descriptions = {
        "Bronze": f"{len(zones['Bronze'])} source таблици. FULL load от GeneraliDWH SQL Server.",
        "Silver": f"{len(zones['Silver'])} таблици. SCD1 Merge (SilverRaw + SilverStg) + Overwrite.",
        "Gold":   f"{len(zones['Gold'])} таблици. SCD2 dim_ + Overwrite fact_. L1→L5 + OW-L1/L2.",
    }

    for r, (zone, items) in enumerate(zones.items(), 5):
        cell_a = ws_index.cell(row=r, column=1, value=zone)
        cell_a.font = XS.zone_link_font()
        cell_a.hyperlink = f"#'{zone_sheet_names[zone]}'!A1"

        ws_index.cell(row=r, column=2, value=zone_descriptions[zone]).font = XS.data_font()
        cnt_cell = ws_index.cell(row=r, column=3, value=len(items))
        cnt_cell.font = Font(name='Calibri', size=10, bold=True, color=Brand.DARK_BLUE_HEX)
        _xl_apply_data_row(ws_index, r, 3, (r - 5) % 2 == 0)

    total = sum(len(v) for v in zones.values())
    ws_index.cell(row=9, column=1, value=f"Total detail sheets: {total}").font = XS.subtitle_font()

    ws_index.merge_cells('A10:C10')
    ws_index['A10'] = (
        "Colour guide:  Teal badge = Direct map  ·  Dark blue = CASE WHEN logic  ·  "
        "Green = COALESCE  ·  Grey = Literal/constant  ·  Light blue = Date formatting"
    )
    ws_index['A10'].font = Font(name='Calibri', size=8, color=Brand.GRAY_HEAD_HEX)
    ws_index['A10'].fill = PatternFill('solid', fgColor=Brand.WARM_BG_HEX)

    ws_index.column_dimensions['A'].width = 20
    ws_index.column_dimensions['B'].width = 65
    ws_index.column_dimensions['C'].width = 10

    # ── Zone Index sheets + Detail sheets ─────────────────────────────────────
    detail_sheets_created = {}

    for zone, items in zones.items():
        zone_sn = zone_sheet_names[zone]
        ws_zone = wb.create_sheet(zone_sn)

        ws_zone.cell(row=1, column=1, value="↖ INDEX").font = XS.nav_font()
        ws_zone['A1'].hyperlink = "#INDEX!A1"

        ws_zone.cell(row=3, column=1, value=f"{zone} Layer — {zone_sn}").font = Font(
            name='Calibri Light', size=16, bold=True, color=Brand.TEAL_HEX
        )
        ws_zone.cell(row=4, column=1, value=f"{len(items)} таблици. Кликни върху таблица за field-level lineage.").font = XS.subtitle_font()

        if zone == "Bronze":
            headers = ["Target Table", "Source Schema", "Source Table", "# Fields", "Load", "Active"]
            col_widths = [45, 15, 40, 10, 10, 8]
        else:
            headers = ["Target Table", "Mode", "Level", "Merge Key", "Source Tables (top 2)", "# Fields"]
            col_widths = [40, 15, 8, 25, 55, 10]

        for c, h in enumerate(headers, 1):
            ws_zone.cell(row=6, column=c, value=h)
        _xl_apply_header_row(ws_zone, 6, len(headers))

        for c, w in enumerate(col_widths, 1):
            ws_zone.column_dimensions[get_column_letter(c)].width = w

        for r_idx, item in enumerate(sorted(items, key=lambda x: x.get("target_table", ""))):
            row = 7 + r_idx
            tbl = item.get("target_table", "")

            prefix = {"Bronze": "B_", "Silver": "SR_", "Gold": "G_"}[zone]
            if zone == "Silver" and item.get("lineage", {}).get("_meta", {}).get("mode") == "overwrite":
                if tbl.startswith("stg_"):
                    prefix = "SS_"
            detail_sn = _xl_safe_sheet_name(tbl, prefix)

            cell_a = ws_zone.cell(row=row, column=1, value=tbl)
            cell_a.font = Font(name='Courier New', size=9, color=Brand.TEAL_HEX, underline='single')
            cell_a.hyperlink = f"#'{detail_sn}'!A1"

            if zone == "Bronze":
                ws_zone.cell(row=row, column=2, value=item.get("source_schema", "")).font = XS.data_font()
                ws_zone.cell(row=row, column=3, value=item.get("source_table", "")).font = XS.code_font()
                ws_zone.cell(row=row, column=4, value=item.get("fields_count", "")).font = XS.data_font()
                ws_zone.cell(row=row, column=5, value="FULL").font = XS.data_font()
                active = "✓" if item.get("is_active", "true").lower() == "true" else "✗"
                ws_zone.cell(row=row, column=6, value=active).font = XS.data_font()
            else:
                ws_zone.cell(row=row, column=2, value=item.get("mode", "").upper()).font = XS.data_font()
                ws_zone.cell(row=row, column=3, value=item.get("level", "")).font = XS.data_font()
                ws_zone.cell(row=row, column=4, value=item.get("merge_key", "")).font = XS.code_font()
                srcs = item.get("source_tables", [])
                ws_zone.cell(row=row, column=5, value="  ·  ".join(srcs[:2])).font = XS.code_font()
                ws_zone.cell(row=row, column=6, value=item.get("fields_count", 0)).font = XS.data_font()

            _xl_apply_data_row(ws_zone, row, len(headers), r_idx % 2 == 0)

            # ── Create detail sheet ───────────────────────────────────────────
            ws_det = wb.create_sheet(detail_sn)
            detail_sheets_created[tbl] = detail_sn

            _xl_add_nav_row(ws_det, 1, [(f"← {zone} Index", zone_sn), ("↖ INDEX", "INDEX")])

            ws_det.cell(row=3, column=1, value=tbl).font = Font(
                name='Calibri Light', size=14, bold=True, color=Brand.DARK_BLUE_HEX
            )

            if zone == "Bronze":
                ws_det.cell(row=4, column=1, value=f"Zone: Bronze (GenDWH_Bronze_LH)").font = XS.subtitle_font()
                src_info = f"Sources: GeneraliDWH SQL Server  ·  {item.get('source_schema','')}.{item.get('source_table','')}"
                ws_det.cell(row=5, column=1, value=src_info).font = XS.subtitle_font()

                cols_str = str(item.get("fields_count", "*"))
                if cols_str == "*":
                    ws_det.cell(row=7, column=1, value="All columns (SELECT *)").font = XS.data_font()
                else:
                    for c, h in enumerate(["Target Field", "Source Column", "Transformation"], 1):
                        ws_det.cell(row=7, column=c, value=h)
                    _xl_apply_header_row(ws_det, 7, 3)

            else:
                lineage = item.get("lineage", {})
                meta = lineage.get("_meta", {})
                mode_str = "SCD2 Merge" if meta.get("mode") == "merge" and zone == "Gold" else \
                           "SCD1 Merge" if meta.get("mode") == "merge" else "Overwrite"
                sk = f"SK={tbl.split('_')[0]}_{tbl.split('_')[1]}_sk" if zone == "Gold" and meta.get("mode") == "merge" else ""

                ws_det.cell(row=4, column=1,
                    value=f"Zone: {zone}  ·  {zone} {meta.get('level','')}  ·  {mode_str}  ·  {sk}  ·  has_current={meta.get('has_current','true')}"
                ).font = XS.subtitle_font()

                srcs = lineage.get("source_tables", [])
                ws_det.cell(row=5, column=1, value=f"Sources: {'  ·  '.join(srcs)}").font = XS.subtitle_font()

                ws_det.cell(row=6, column=1,
                    value="Legend:  Direct map = field copied as-is  ·  Source Table '—' = literal/constant  ·  Transformation shows SQL logic"
                ).font = Font(name='Calibri', size=8, color=Brand.GRAY_HEAD_HEX)

                # Field-level table
                if zone == "Gold" and meta.get("mode") == "merge":
                    det_headers = ["Target Field", "Data Type", "Transformation", "Source Table", "Source Column", "Expression (SQL)", "Join Key", "SCD2 Tracked"]
                    det_widths = [30, 12, 14, 40, 25, 50, 20, 12]
                else:
                    det_headers = ["Target Field", "Data Type", "Transformation", "Source Table", "Source Column", "Expression (SQL)", "Join Key"]
                    det_widths = [30, 12, 14, 40, 25, 50, 20]

                hdr_row = 7
                for c, h in enumerate(det_headers, 1):
                    ws_det.cell(row=hdr_row, column=c, value=h)
                _xl_apply_header_row(ws_det, hdr_row, len(det_headers))

                for c, w in enumerate(det_widths, 1):
                    ws_det.column_dimensions[get_column_letter(c)].width = w

                for f_idx, field in enumerate(lineage.get("fields", [])):
                    fr = hdr_row + 1 + f_idx
                    is_alt = f_idx % 2 == 0
                    tt = field.get("transformation_type", "")

                    ws_det.cell(row=fr, column=1, value=field.get("target_field", "")).font = XS.code_font(Brand.DARK_BLUE_HEX)
                    ws_det.cell(row=fr, column=2, value=field.get("data_type", "")).font = XS.meta_font()

                    badge_cell = ws_det.cell(row=fr, column=3, value=tt.replace("_", " ").title() if tt else "")
                    badge_cell.font = XS.badge_font()
                    badge_cell.fill = XS.badge_fill(tt)
                    badge_cell.alignment = Alignment(horizontal='center')

                    ws_det.cell(row=fr, column=4, value=field.get("source_table", "")).font = XS.code_font(Brand.DARK_BLUE_HEX)
                    ws_det.cell(row=fr, column=5, value=field.get("source_column", "")).font = XS.code_font(Brand.TEAL_HEX)

                    expr = field.get("expression", "")
                    ws_det.cell(row=fr, column=6, value=expr).font = Font(name='Courier New', size=8, color=Brand.DARK_HEX)

                    ws_det.cell(row=fr, column=7, value=field.get("join_key", "")).font = XS.data_font()

                    if len(det_headers) == 8:
                        ws_det.cell(row=fr, column=8, value="✓ SCD2" if field.get("join_key") else "").font = Font(
                            name='Calibri', size=9, bold=True, color=Brand.DARK_BLUE_HEX
                        )

                    _xl_apply_data_row(ws_det, fr, len(det_headers), is_alt)
                    # Preserve badge fill on column 3
                    ws_det.cell(row=fr, column=3).fill = XS.badge_fill(tt)
                    ws_det.cell(row=fr, column=3).font = XS.badge_font()

                # Freeze panes
                ws_det.freeze_panes = f"A{hdr_row + 1}"

    # Save
    out_path = f"{PATHS['excel_docs']}/GenDWH_Data_Lineage.xlsx"
    wb.save(out_path)
    print(f"   ✓ Saved: {out_path}")
    print(f"   📊 {len(detail_sheets_created)} detail sheets + 3 zone indexes + INDEX")
    return out_path


# ── Data Dictionary Excel ─────────────────────────────────────────────────────

def generate_dictionary_excel(lh_tables=None):
    """Generate Data Dictionary Excel with table & column-level detail."""
    print("\n📗 Generating Data Dictionary Excel...")
    wb = openpyxl.Workbook()
    ws_index = wb.active
    ws_index.title = "INDEX"
    now_str = datetime.now().strftime("%B %Y")

    if lh_tables is None:
        rows = spark.sql("SELECT * FROM doc_lakehouse_tables").collect()
        lh_tables = [r.asDict() for r in rows]

    # INDEX sheet
    ws_index.merge_cells('A1:D1')
    ws_index['A1'] = "GenDWH — Data Dictionary"
    ws_index['A1'].font = XS.title_font()
    ws_index.merge_cells('A2:D2')
    ws_index['A2'] = f"Table & column descriptions by Lakehouse  ·  {now_str}"
    ws_index['A2'].font = XS.subtitle_font()

    headers = ["Lakehouse", "Table", "Columns", "Rows"]
    for c, h in enumerate(headers, 1):
        ws_index.cell(row=4, column=c, value=h)
    _xl_apply_header_row(ws_index, 4, 4)

    ws_index.column_dimensions['A'].width = 30
    ws_index.column_dimensions['B'].width = 45
    ws_index.column_dimensions['C'].width = 10
    ws_index.column_dimensions['D'].width = 15

    for r_idx, tbl in enumerate(sorted(lh_tables, key=lambda x: (x.get("lakehouse",""), x.get("table_name","")))):
        row = 5 + r_idx
        tbl_name = tbl.get("table_name", "")
        det_sn = _xl_safe_sheet_name(tbl_name, "T_")

        ws_index.cell(row=row, column=1, value=tbl.get("lakehouse", "")).font = XS.data_font()

        cell_b = ws_index.cell(row=row, column=2, value=tbl_name)
        cell_b.font = Font(name='Courier New', size=9, color=Brand.TEAL_HEX, underline='single')
        cell_b.hyperlink = f"#'{det_sn}'!A1"

        ws_index.cell(row=row, column=3, value=tbl.get("column_count", 0)).font = XS.data_font()
        ws_index.cell(row=row, column=4, value=tbl.get("row_count", -1)).font = XS.data_font()
        _xl_apply_data_row(ws_index, row, 4, r_idx % 2 == 0)

        # Detail sheet
        ws_det = wb.create_sheet(det_sn)
        _xl_add_nav_row(ws_det, 1, [("↖ INDEX", "INDEX")])
        ws_det.cell(row=3, column=1, value=tbl_name).font = Font(
            name='Calibri Light', size=14, bold=True, color=Brand.DARK_BLUE_HEX
        )
        ws_det.cell(row=4, column=1,
            value=f"Lakehouse: {tbl.get('lakehouse','')}  ·  Rows: {tbl.get('row_count', '?')}  ·  Columns: {tbl.get('column_count', '?')}"
        ).font = XS.subtitle_font()

        col_headers = ["Column Name", "Data Type"]
        for c, h in enumerate(col_headers, 1):
            ws_det.cell(row=6, column=c, value=h)
        _xl_apply_header_row(ws_det, 6, 2)
        ws_det.column_dimensions['A'].width = 40
        ws_det.column_dimensions['B'].width = 25

        schema = json.loads(tbl.get("schema_json", "[]"))
        for s_idx, col_info in enumerate(schema):
            sr = 7 + s_idx
            ws_det.cell(row=sr, column=1, value=col_info.get("name", "")).font = XS.code_font()
            ws_det.cell(row=sr, column=2, value=col_info.get("type", "")).font = XS.meta_font()
            _xl_apply_data_row(ws_det, sr, 2, s_idx % 2 == 0)

        ws_det.freeze_panes = "A7"

    out_path = f"{PATHS['excel_docs']}/GenDWH_Data_Dictionary.xlsx"
    wb.save(out_path)
    print(f"   ✓ Saved: {out_path}")
    return out_path


# ── Workspace Inventory Excel ─────────────────────────────────────────────────

def generate_inventory_excel(inventory=None):
    """Generate Workspace Inventory Excel."""
    print("\n📗 Generating Workspace Inventory Excel...")
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Inventory"
    now_str = datetime.now().strftime("%B %Y")

    if inventory is None:
        rows = spark.sql("SELECT * FROM doc_workspace_inventory").collect()
        inventory = [r.asDict() for r in rows]

    ws.merge_cells('A1:F1')
    ws['A1'] = "GenDWH — Workspace Inventory"
    ws['A1'].font = XS.title_font()
    ws.merge_cells('A2:F2')
    ws['A2'] = f"All accessible Fabric workspaces and items  ·  {now_str}"
    ws['A2'].font = XS.subtitle_font()

    headers = ["Workspace", "Item Type", "Item Name", "Item ID", "Description"]
    for c, h in enumerate(headers, 1):
        ws.cell(row=4, column=c, value=h)
    _xl_apply_header_row(ws, 4, len(headers))

    col_widths = [35, 20, 45, 40, 50]
    for c, w in enumerate(col_widths, 1):
        ws.column_dimensions[get_column_letter(c)].width = w

    for r_idx, item in enumerate(sorted(inventory, key=lambda x: (x.get("workspace_name",""), x.get("item_type",""), x.get("item_name","")))):
        row = 5 + r_idx
        ws.cell(row=row, column=1, value=item.get("workspace_name", "")).font = XS.data_font()
        ws.cell(row=row, column=2, value=item.get("item_type", "")).font = XS.data_font()
        ws.cell(row=row, column=3, value=item.get("item_name", "")).font = XS.code_font()
        ws.cell(row=row, column=4, value=item.get("item_id", "")).font = Font(name='Courier New', size=8, color=Brand.GRAY_HEAD_HEX)
        ws.cell(row=row, column=5, value=item.get("item_description", "")).font = XS.data_font()
        _xl_apply_data_row(ws, row, len(headers), r_idx % 2 == 0)

    ws.freeze_panes = "A5"
    ws.auto_filter.ref = f"A4:E{4 + len(inventory)}"

    # Summary sheet
    ws_sum = wb.create_sheet("Summary", 0)
    ws_sum.merge_cells('A1:C1')
    ws_sum['A1'] = "Workspace Summary"
    ws_sum['A1'].font = XS.title_font()

    # Count by workspace × type
    from collections import Counter
    ws_types = Counter((i["workspace_name"], i["item_type"]) for i in inventory)
    ws_names = sorted(set(i["workspace_name"] for i in inventory))
    item_types = sorted(set(i["item_type"] for i in inventory))

    for c, t in enumerate(item_types, 2):
        ws_sum.cell(row=3, column=c, value=t)
    _xl_apply_header_row(ws_sum, 3, len(item_types) + 1)
    ws_sum.cell(row=3, column=1, value="Workspace")

    for r, ws_name in enumerate(ws_names, 4):
        ws_sum.cell(row=r, column=1, value=ws_name).font = Font(name='Calibri', size=10, bold=True, color=Brand.TEAL_HEX)
        for c, t in enumerate(item_types, 2):
            cnt = ws_types.get((ws_name, t), 0)
            if cnt:
                ws_sum.cell(row=r, column=c, value=cnt).font = XS.data_font()
        _xl_apply_data_row(ws_sum, r, len(item_types) + 1, (r - 4) % 2 == 0)

    for c in range(1, len(item_types) + 2):
        ws_sum.column_dimensions[get_column_letter(c)].width = 22

    out_path = f"{PATHS['excel_docs']}/GenDWH_Workspace_Inventory.xlsx"
    wb.save(out_path)
    print(f"   ✓ Saved: {out_path}")
    return out_path



def _generate_glossary_via_ai(ai_results, api_key=None):
    """Ask Claude to generate/enrich a business glossary from lineage data."""
    terms = set()
    for tbl, lineage in (ai_results or {}).items():
        for field in lineage.get("fields", []):
            bl = field.get("business_logic", "")
            if bl:
                terms.add(bl[:100])

    sample_tables = list((ai_results or {}).keys())[:20]
    prompt = f"""Based on the GenDWH data warehouse platform (Generali insurance, Microsoft Fabric, IFRS 17),
generate a comprehensive business glossary in JSON format.

The platform contains these tables (sample): {json.dumps(sample_tables)}

Generate glossary entries covering:
1. Insurance concepts (policy, claim, premium, deductible, etc.)
2. IFRS 17 concepts (UoA, DAC, UPR, LRC, PAA, BSP DAC, cohort, etc.)
3. Ceded reinsurance (ceded/assumed, CCLPRP, CCLRES, CPRC, CRECPOS, FPSL, etc.)
4. DWH technical concepts (SCD1, SCD2, surrogate key, medallion architecture, etc.)
5. Agent and sales channel concepts

Format: JSON array of objects:
[{{"term_bg": "...", "term_en": "...", "category": "Insurance|IFRS17|Reinsurance|DWH|Sales", "definition": "..."}}]

Include 40-60 terms. Respond with ONLY valid JSON."""

    try:
        response = call_claude(prompt, api_key=api_key)
        clean = response.strip()
        if clean.startswith("```"):
            clean = re.sub(r'^```\w*\n?', '', clean)
            clean = re.sub(r'\n?```$', '', clean)
        return json.loads(clean)
    except Exception as e:
        print(f"   ⚠ Glossary AI error: {e}")
        return []


def generate_business_glossary_word(ai_results=None, api_key=None):
    """Generate Business Glossary Word document."""
    print("\n📘 Generating Business Glossary Word...")
    now_str = datetime.now().strftime("%B %Y")

    doc = DocxDocument()
    docx_setup_styles(doc)
    docx_add_header_footer(doc, "Business Glossary | GenDWH")

    docx_title_page(doc,
        product="GENDWH",
        title="Business Glossary",
        subtitle="Речник на бизнес и технически термини",
        meta_line="GenDWH Platform  |  Microsoft Fabric  |  InspirIT",
        date_str=now_str,
    )

    docx_version_table(doc, [{
        "version": "1.0",
        "date": datetime.now().strftime("%Y-%m-%d"),
        "author": "InspirIT AI",
        "description": "Auto-generated from lineage analysis",
    }])

    terms = _generate_glossary_via_ai(ai_results, api_key)
    if not terms:
        terms = []

    categories = {}
    for t in terms:
        cat = t.get("category", "Other")
        categories.setdefault(cat, []).append(t)

    for cat in sorted(categories.keys()):
        doc.add_heading(cat, level=1)

        headers = ["Термин (BG)", "Term (EN)", "Дефиниция"]
        rows = [[t.get("term_bg",""), t.get("term_en",""), t.get("definition","")] for t in categories[cat]]
        docx_add_table(doc, headers, rows, col_widths_cm=[4, 4, 10])
        doc.add_paragraph()

    out_path = f"{PATHS['word_docs']}/GenDWH_Business_Glossary.docx"
    doc.save(out_path)
    print(f"   ✓ Saved: {out_path}")
    return out_path


def generate_system_architecture_word(inventory=None, activities=None):
    """Generate System Architecture Word document from pipeline definitions."""
    print("\n📘 Generating System Architecture Word...")
    now_str = datetime.now().strftime("%B %Y")

    doc = DocxDocument()
    docx_setup_styles(doc)
    docx_add_header_footer(doc, "System Architecture | GenDWH")

    docx_title_page(doc,
        product="GENDWH",
        title="System Architecture",
        subtitle="Technical Documentation",
        meta_line="GenDWH Platform  |  Microsoft Fabric  |  InspirIT",
        date_str=now_str,
    )

    docx_version_table(doc, [{
        "version": "1.0",
        "date": datetime.now().strftime("%Y-%m-%d"),
        "author": "InspirIT AI",
        "description": "Auto-generated from Fabric workspace discovery",
    }])

    # Section 1: Platform Overview
    doc.add_heading("Platform Overview", level=1)

    if inventory is None:
        try:
            rows = spark.sql("SELECT * FROM doc_workspace_inventory").collect()
            inventory = [r.asDict() for r in rows]
        except:
            inventory = []

    ws_names = sorted(set(i.get("workspace_name","") for i in inventory))
    doc.add_paragraph(
        f"The GenDWH platform is deployed across {len(ws_names)} Microsoft Fabric workspace(s). "
        f"The solution implements a medallion architecture with Bronze, Silver, and Gold layers, "
        f"orchestrated by metadata-driven pipelines."
    )

    if ws_names:
        doc.add_heading("Workspaces", level=2)
        for ws in ws_names:
            items = [i for i in inventory if i.get("workspace_name") == ws]
            type_counts = {}
            for i in items:
                t = i.get("item_type", "")
                type_counts[t] = type_counts.get(t, 0) + 1

            p = doc.add_paragraph()
            run = p.add_run(ws)
            run.font.bold = True
            run.font.color.rgb = Brand.RED
            desc_parts = [f"{c} {t}" for t, c in sorted(type_counts.items())]
            p.add_run(f" — {', '.join(desc_parts)}")

    # Section 2: Pipeline Architecture
    doc.add_heading("Pipeline Architecture", level=1)

    if activities is None:
        try:
            rows = spark.sql("SELECT * FROM doc_pipeline_activities").collect()
            activities = [r.asDict() for r in rows]
        except:
            activities = []

    pipeline_names = sorted(set(a.get("pipeline_name","") for a in activities))
    for pname in pipeline_names:
        doc.add_heading(pname, level=2)
        pipe_acts = [a for a in activities if a.get("pipeline_name") == pname]

        doc.add_paragraph(f"Pipeline with {len(pipe_acts)} activities.")

        headers = ["Activity", "Type", "State", "Description"]
        rows_data = [
            [a.get("activity_name",""), a.get("activity_type",""),
             a.get("state",""), a.get("description","")[:60]]
            for a in sorted(pipe_acts, key=lambda x: x.get("activity_name",""))
        ]
        docx_add_table(doc, headers, rows_data, col_widths_cm=[5, 3, 2, 8])
        doc.add_paragraph()

        # Dependency map
        doc.add_heading("Dependency Map", level=3)
        headers = ["Activity", "Depends On", "Condition"]
        dep_rows = []
        for a in pipe_acts:
            deps = json.loads(a.get("depends_on", "[]"))
            if deps:
                for d in deps:
                    dep_rows.append([
                        a.get("activity_name",""),
                        d.get("activity",""),
                        ", ".join(d.get("conditions", []))
                    ])
        if dep_rows:
            docx_add_table(doc, headers, dep_rows, col_widths_cm=[5, 5, 3])

    # Section 3: Medallion Architecture
    doc.add_heading("Medallion Architecture", level=1)
    doc.add_paragraph(
        "GenDWH follows an extended medallion architecture with the following layers:"
    )

    layers = [
        ("Bronze", "Full copy from GeneraliDWH SQL Server. Raw data, no transformations. Overwrite schema on each load."),
        ("Silver Raw", "Homogenized nomenclatures (nom_) and raw transactional tables (raw_). SCD Type 1 merge — current values only."),
        ("Silver Staging", "IFRS calculations, aggregates, and staging tables (stg_). Prepares data for Gold fact tables."),
        ("Gold DWH", "Classic dimensional model. SCD Type 2 for dimensions (surrogate keys, is_current, effective_date). Fact tables with FK references."),
        ("Platinum", "Underwriting calculations. Async trigger from Gold. Separate pipeline."),
    ]
    headers = ["Layer", "Description"]
    docx_add_table(doc, headers, layers, col_widths_cm=[3, 15])

    out_path = f"{PATHS['word_docs']}/GenDWH_System_Architecture.docx"
    doc.save(out_path)
    print(f"   ✓ Saved: {out_path}")
    return out_path


def generate_tech_specs_word(activities=None, queries=None):
    """Generate Technical Specifications Word document."""
    print("\n📘 Generating Technical Specifications Word...")
    now_str = datetime.now().strftime("%B %Y")

    doc = DocxDocument()
    docx_setup_styles(doc)
    docx_add_header_footer(doc, "Technical Specifications | GenDWH")

    docx_title_page(doc,
        product="GENDWH",
        title="Technical Specifications",
        subtitle="Variables, Error Handling, Metadata Tables",
        meta_line="GenDWH Platform  |  Microsoft Fabric  |  InspirIT",
        date_str=now_str,
    )

    docx_version_table(doc, [{
        "version": "1.0",
        "date": datetime.now().strftime("%Y-%m-%d"),
        "author": "InspirIT AI",
        "description": "Auto-generated from pipeline definitions and metadata",
    }])

    # Section 1: Metadata Tables
    doc.add_heading("Metadata Tables", level=1)
    doc.add_paragraph(
        "All metadata tables reside in GenDWH_Administration_LH. "
        "The Orchestrator Notebook reads the appropriate table based on the meta_table parameter passed by the pipeline."
    )

    if queries is None:
        try:
            rows = spark.sql("SELECT * FROM doc_metadata_queries").collect()
            queries = [r.asDict() for r in rows]
        except:
            queries = []

    meta_tables = sorted(set((q.get("meta_table",""), q.get("layer",""), q.get("level",""), q.get("mode","")) for q in queries))
    if meta_tables:
        headers = ["Metadata Table", "Layer", "Level", "Mode"]
        rows_data = [list(m) for m in meta_tables]
        docx_add_table(doc, headers, rows_data, col_widths_cm=[8, 2, 2, 3])

    # Section 2: Pipeline Variables
    doc.add_heading("Pipeline Variables", level=1)
    doc.add_paragraph(
        "Library variables are managed centrally through GenDWH_Variables library. "
        "The pipeline accesses them as @pipeline().libraryVariables.<name> — no hardcoded IDs anywhere."
    )

    # Section 3: Error Handling
    doc.add_heading("Error Handling", level=1)
    doc.add_paragraph(
        "Every critical activity has a dedicated error chain: on failure, an Office 365 Outlook notification is sent, "
        "followed by a Fail activity that terminates the pipeline with error code 500. "
        "This ensures the pipeline never fails silently."
    )

    # Section 4: Schedule & Retry
    doc.add_heading("Schedule & Retry Policy", level=1)

    schedule_data = [
        ["Schedule", "Mon–Fri, 03:00"],
        ["Timezone", "FLE Standard Time"],
        ["Retry Policy", "5 retries / 30s interval"],
        ["Layers", "Bronze → Silver → Gold → Platinum (async)"],
    ]
    headers = ["Parameter", "Value"]
    docx_add_table(doc, headers, schedule_data, col_widths_cm=[5, 13])

    # Section 5: Design Decisions
    doc.add_heading("Key Design Decisions", level=1)

    decisions = [
        ("Metadata-Driven Pipeline", "No hardcoded table names in pipeline activities. All definitions read from gen_adm_* tables at runtime. Adding a new table requires only an INSERT into metadata."),
        ("Single Orchestrator Notebook", "All Silver and Gold transformation activities call the same Orchestrator_NB_ID notebook, parameterized by meta_table, mode, load_type, and layer."),
        ("SCD1 vs SCD2 Strategy", "Silver uses SCD1 (current values only). Gold merge levels use SCD2 (full history with effective dates). Gold overwrite levels use SCD1 where history is not needed."),
        ("Async Platinum Trigger", "Platinum pipeline starts with waitOnCompletion=false. The orchestrating pipeline reports as completed immediately after firing the trigger."),
    ]
    for title, desc in decisions:
        doc.add_heading(title, level=2)
        doc.add_paragraph(desc)

    out_path = f"{PATHS['word_docs']}/GenDWH_Technical_Specifications.docx"
    doc.save(out_path)
    print(f"   ✓ Saved: {out_path}")
    return out_path


def generate_lineage_overview_word(ai_results=None):
    """Generate Data Lineage Overview Word document — narrative flow descriptions."""
    print("\n📘 Generating Data Lineage Overview Word...")
    now_str = datetime.now().strftime("%B %Y")

    doc = DocxDocument()
    docx_setup_styles(doc)
    docx_add_header_footer(doc, "Data Lineage | GenDWH")

    docx_title_page(doc,
        product="GENDWH",
        title="Data Lineage",
        subtitle="Проследяване на данните Bronze → Silver → Gold",
        meta_line="GenDWH Platform  |  Microsoft Fabric  |  InspirIT",
        date_str=now_str,
    )

    docx_version_table(doc, [{
        "version": "1.0",
        "date": datetime.now().strftime("%Y-%m-%d"),
        "author": "InspirIT AI",
        "description": "Auto-generated from AI lineage analysis",
    }])

    if ai_results is None:
        ai_results = _load_cache()

    doc.add_heading("Overview", level=1)
    doc.add_paragraph(
        "This document traces data flows from GeneraliDWH SQL Server (Bronze) "
        "through SilverRaw/SilverStg to Gold DWH tables. For each flow, source tables, "
        "transformation logic, and Gold targets are described."
    )

    # Group tables by layer and level
    layers = {"Silver": {}, "Gold": {}}
    for tbl, lineage in (ai_results or {}).items():
        meta = lineage.get("_meta", {})
        layer = meta.get("layer", "")
        level = meta.get("level", "")
        if layer in layers:
            layers[layer].setdefault(level, []).append((tbl, lineage))

    for layer in ["Silver", "Gold"]:
        doc.add_heading(f"{layer} Layer", level=1)
        for level in sorted(layers[layer].keys()):
            doc.add_heading(f"{layer} {level}", level=2)
            for tbl, lineage in sorted(layers[layer][level], key=lambda x: x[0]):
                doc.add_heading(tbl, level=3)
                meta = lineage.get("_meta", {})

                srcs = lineage.get("source_tables", [])
                if srcs:
                    doc.add_paragraph(f"Sources: {', '.join(srcs)}")

                cte = lineage.get("cte_chain", [])
                if cte:
                    doc.add_paragraph(f"CTE chain: {' → '.join(cte)}")

                fields = lineage.get("fields", [])
                complex_fields = [f for f in fields if f.get("transformation_type") not in ("direct_map", "literal", "")]
                if complex_fields:
                    p = doc.add_paragraph("Key transformations: ")
                    for f in complex_fields[:5]:
                        doc.add_paragraph(
                            f"  • {f.get('target_field','')}: {f.get('business_logic', f.get('expression','')[:80])}",
                            style='List Bullet'
                        )

    out_path = f"{PATHS['word_docs']}/GenDWH_Data_Lineage_Overview.docx"
    doc.save(out_path)
    print(f"   ✓ Saved: {out_path}")
    return out_path


def phase_4_doc_generation(ai_results=None, inventory=None, activities=None, queries=None, lh_tables=None, bronze_meta=None):
    """Phase 4 — Generate all Excel and Word documents."""
    print("═" * 60)
    print("PHASE 4 — DOCUMENT GENERATION")
    print("═" * 60)

    api_key = get_claude_api_key()
    outputs = {}

    outputs["lineage_xlsx"]      = generate_lineage_excel(ai_results, bronze_meta, queries)
    outputs["dictionary_xlsx"]   = generate_dictionary_excel(lh_tables)
    outputs["inventory_xlsx"]    = generate_inventory_excel(inventory)
    outputs["glossary_docx"]     = generate_business_glossary_word(ai_results, api_key)
    outputs["architecture_docx"] = generate_system_architecture_word(inventory, activities)
    outputs["tech_specs_docx"]   = generate_tech_specs_word(activities, queries)
    outputs["lineage_docx"]      = generate_lineage_overview_word(ai_results)

    print(f"\n── Generated Documents ──")
    for key, path in outputs.items():
        size = os.path.getsize(path) if os.path.exists(path) else 0
        print(f"   {key:30s} {size/1024:.0f} KB  {path}")

    print("\n✓ Phase 4 complete")
    return outputs



# CELL 9 ── Phase 5 — Versioning ───

import shutil


def _compute_run_fingerprint(inventory, queries, ai_results):
    """Compute a fingerprint of the current state for change detection."""
    parts = []
    for item in sorted(inventory or [], key=lambda x: x.get("item_id","")):
        parts.append(f"{item.get('item_id','')}:{item.get('item_name','')}")
    for q in sorted(queries or [], key=lambda x: x.get("target_table","")):
        parts.append(f"{q.get('target_table','')}:{q.get('query_hash','')}")
    return sha256("|".join(parts))


def _get_previous_version():
    """Get the latest version record from Delta table."""
    try:
        rows = spark.sql(
            "SELECT * FROM doc_version_history ORDER BY run_timestamp DESC LIMIT 1"
        ).collect()
        if rows:
            return rows[0].asDict()
    except:
        pass
    return None


def phase_5_versioning(inventory=None, queries=None, ai_results=None, doc_outputs=None):
    """
    Phase 5 — Version tracking and archival.
    """
    print("═" * 60)
    print("PHASE 5 — VERSIONING")
    print("═" * 60)

    run_ts = datetime.now(timezone.utc)
    run_ts_str = run_ts.isoformat()
    fingerprint = _compute_run_fingerprint(inventory, queries, ai_results)

    prev = _get_previous_version()
    prev_fingerprint = prev.get("fingerprint", "") if prev else ""
    prev_version = prev.get("version", "0.0") if prev else "0.0"

    if fingerprint == prev_fingerprint:
        print("   ℹ No changes detected since last run")
        new_version = prev_version
        change_type = "no_change"
    else:
        major, minor = prev_version.split(".")
        new_version = f"{major}.{int(minor) + 1}"
        change_type = "updated"
        print(f"   🔄 Changes detected! Version: {prev_version} → {new_version}")

    version_dir = f"{PATHS['versions']}/v{new_version}_{run_ts.strftime('%Y%m%d_%H%M%S')}"
    os.makedirs(version_dir, exist_ok=True)

    if doc_outputs:
        for key, src_path in doc_outputs.items():
            if os.path.exists(src_path):
                dest = os.path.join(version_dir, os.path.basename(src_path))
                shutil.copy2(src_path, dest)

    print(f"   📁 Archived to: {version_dir}")

    version_record = {
        "version":          new_version,
        "run_timestamp":    run_ts_str,
        "fingerprint":      fingerprint,
        "change_type":      change_type,
        "prev_version":     prev_version,
        "workspace_count":  len(set(i.get("workspace_id","") for i in (inventory or []))),
        "item_count":       len(inventory or []),
        "query_count":      len(queries or []),
        "field_count":      sum(len(l.get("fields",[])) for l in (ai_results or {}).values()),
        "documents":        json.dumps(list((doc_outputs or {}).keys())),
        "archive_path":     version_dir,
    }

    df = spark.createDataFrame([version_record])

    try:
        df.write.format("delta").mode("append").saveAsTable("doc_version_history")
    except:
        df.write.format("delta").mode("overwrite").saveAsTable("doc_version_history")

    print(f"\n── Version Summary ──")
    print(f"   Version:     {new_version}")
    print(f"   Timestamp:   {run_ts_str}")
    print(f"   Change type: {change_type}")
    print(f"   Fingerprint: {fingerprint}")
    print(f"   Workspaces:  {version_record['workspace_count']}")
    print(f"   Items:       {version_record['item_count']}")
    print(f"   SQL queries: {version_record['query_count']}")
    print(f"   Fields:      {version_record['field_count']}")
    print(f"   Documents:   {len(doc_outputs or {})}")

    print("\n✓ Phase 5 complete")
    return version_record



# CELL 10 ── Main Orchestrator ───

def run_all(force_ai_rerun=False):
    """
    Run the complete documentation pipeline: Discovery → Metadata → AI →
    Knowledge Build → Doc Generation → Versioning.

    Args:
        force_ai_rerun: If True, re-analyze all SQL queries even if cached
    """
    start = time.time()
    print("╔══════════════════════════════════════════════════════════╗")
    print("║  GenDWH — Automated Documentation System                ║")
    print("║  Full Run                                               ║")
    print(f"║  {datetime.now().strftime('%Y-%m-%d %H:%M:%S'):56s} ║")
    print("╚══════════════════════════════════════════════════════════╝")

    # Phase 0 — Discovery
    inventory = phase_0_discovery()

    # Phase 1 — Metadata Extraction
    meta = phase_1_metadata_extraction(inventory)
    activities  = meta["activities"]
    lh_tables   = meta["tables"]
    queries     = meta["queries"]
    bronze_meta = meta["bronze_meta"]

    # Phase 2 — AI Analysis
    ai_results = phase_2_ai_analysis(queries, force_rerun=force_ai_rerun)

    # Phase 3 — Knowledge Build (JSONL)
    phase_3_knowledge_build(
        ai_results=ai_results,
        inventory=inventory,
        activities=activities,
        queries=queries,
        lh_tables=lh_tables,
        bronze_meta=bronze_meta,
    )

    # Phase 4 — Document Generation
    doc_outputs = phase_4_doc_generation(
        ai_results=ai_results,
        inventory=inventory,
        activities=activities,
        queries=queries,
        lh_tables=lh_tables,
        bronze_meta=bronze_meta,
    )

    # Phase 5 — Versioning
    version = phase_5_versioning(inventory, queries, ai_results, doc_outputs)

    elapsed = time.time() - start
    print(f"\n{'═' * 60}")
    print(f"✅ ALL PHASES COMPLETE — v{version['version']} — {elapsed:.0f}s")
    print(f"{'═' * 60}")
    print(f"\n📂 Documents available at:")
    print(f"   Excel: {PATHS['excel_docs']}/")
    print(f"   Word:  {PATHS['word_docs']}/")
    print(f"   Archive: {version['archive_path']}")
    print(f"\n📊 Delta tables created:")
    print(f"   doc_workspace_inventory  — All workspace items")
    print(f"   doc_pipeline_activities  — Pipeline activity definitions")
    print(f"   doc_lakehouse_tables     — Table schemas and row counts")
    print(f"   doc_metadata_queries     — SQL queries from metadata")
    print(f"   doc_bronze_metadata      — Bronze layer source tables")
    print(f"   doc_ai_lineage           — Field-level lineage (AI-analyzed)")
    print(f"   doc_version_history      — Version tracking")

    return {
        "version": version,
        "documents": doc_outputs,
        "inventory": inventory,
        "ai_results": ai_results,
    }


print(f"✓ GenDWH Documentation LIB loaded — v{LIB_VERSION}")